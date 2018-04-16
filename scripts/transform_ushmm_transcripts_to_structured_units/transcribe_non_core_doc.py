import sys, glob, os
helper_path = os.path.join("..", "..", "utils")
sys.path.insert(0, helper_path)
import helper_mongo as h

os.chdir("../../data/")
from docx import Document
from subprocess import call

import pprint
import constants
import re
from collections import defaultdict

TRACKER = constants.TRACKER_COLLECTION
OUTPUT = constants.OUTPUT_COLLECTION
DB = constants.DB

pp = pprint.PrettyPrinter(indent=4)

def safePrint(str_):
    return ''.join(i for i in str_ if ord(i) < 128)

def getUnstructured042Units(filename):
    """
    Returns the unstructured units of the RG.50-402 series
    These interviews did not have any indi
    """
    doc = Document(filename)
    units = list()
    # all interviews start with a header
    isHeader = True
    # iterate over all paragraphs to get text units
    for para in doc.paragraphs:
        paragraph = para.text
        
        # ensure paragraph is not just empty line
        hasText = paragraph.lstrip()
        # ensure it is not an empty line
        if hasText:

            if isHeader:
                if 'beep' in paragraph.lower():
                    isHeader = False
            else:
                # marks the end of the interview
                if 'USHMM Archives' in paragraph or "wentworth films" in paragraph.lower() :
                    break
                elif 'beep' not in paragraph.lower():
                    units.append({'unit':paragraph})
    return units

def getUnstructured926Units(filename):
    """
    Returns the unstructured units of the RG.50-402 series
    These interviews did not have any indi
    """
    doc = Document(filename)
    units = list()
    n = re.compile('track [0-9][0-9]')

    # iterate over all paragraphs to get text units
    for para in doc.paragraphs:
        paragraph = para.text
        
        # ensure paragraph is not just empty line
        hasText = paragraph.lstrip()
        # ensure it is not an empty line
        if hasText:
            isHeader = n.match(paragraph.lower())

            # marks the end of the interview
            if "story preservation initiative" in paragraph.lower() or "copyright" in paragraph.lower() :
                break
            elif not isHeader:
                units.append({'unit':paragraph})
    pp.pprint(units)
    return units
def getTextUnits(filename):
    """
    Returns the text units for a given file in the non-core asset
    Uses regex to identify common patterns and uses specific backup methods
    depending on the interview shelfmark series, in case they are highly
    unstructured
    """
    doc = Document(filename)
    units = list()
    
    unit_tracker = defaultdict(int)
    
    non_units = ["name:", "date:", "thesis:", "currently:", "note", "comment", "grandparents:", "transcript:", "note:"]

    
    # iterate over all paragraphs to get text units
    for para in doc.paragraphs:
        paragraph = para.text
        
        # ensure it is not an empty line
        if paragraph:
            # get first word
            unit_type = paragraph.partition(' ')[0]
            # in case it is in the format of e.g 'George Salton:'
            
            # e.g AJ:, B.
            m = re.compile('[A-Z][A-Z]?[.|:|-]')
            type1 = m.match(unit_type)

            # e.g [WJ]
            n = re.compile('\[[A-Z][A-Z]\]')
            typer2= n.match(unit_type)

            # timestamp e.g 15:01:27
            o = re.compile('[0-9]?[0-9]:[0-9][0-9]:[0-9][0-9]')
            type3 = o.match(unit_type)           

            ongoing_answer = ""
            
            # else parse them according to formatting guidelines
            if ("Question:" in unit_type or
                type1 or
                "Answer:" in unit_type or 
                typer2 or
                type3):

                # check if there was an ongoing paragraph
                units.append({'unit': paragraph})
                '''
                if ongoing_answer:                
                    units.append({'unit': ongoing_answer})

                # reset it
                ongoing_answer = ""
                ongoing_answer += paragraph
                '''
                # update tracker
                unit_tracker[unit_type] += 1

            elif (unit_type.endswith(':') and
                    unit_type.lower() not in non_units and
                    unit_type[:-1].isalpha()):
                
                units.append({'unit': paragraph})
                # update tracker
                unit_tracker[unit_type] += 1
            
            # backup method,in case it is in the format of e.g 'George Salton:'
            elif len(paragraph.split()) > 3:
                backup_type = paragraph.split()[1]
                backup_two = paragraph.split()[2]

                if ((':' in backup_type and backup_type.lower() not in non_units) or
                    (':' in backup_two and backup_two.lower() not in non_units)): 
                    units.append({'unit': paragraph})
                    # update tracker
                    unit_tracker[unit_type] += 1

            # if it is an ongoing answer
            """
            elif ongoing_answer:
                if not any(non_unit in paragraph.lower() for non_unit in non_units):
                    ongoing_answer += paragraph
            """
    
    if len(unit_tracker) < 2:
        if "50.042" in filename:
            units = getUnstructured042Units(filename)
        elif "50.926" in filename:
            print(filename)
            units = getUnstructured926Units(filename)
        else:
          return []
    
    return units

def createStructuredTranscriptDocx():
    """
    Processes the 509 doc files beloging to the core asset in data
    Core asset is identified by numbers RG-50.030, RG-50.106, RG-50.549
    """
    core_doc_asset = []
    missing_count = 0
    # get all the docx files that are part of the core asset
    for file in glob.glob("*.doc"):

        # RG numbers for the core asset
        if ("RG-50.030" not in file and
            "RG-50.106" not in file and
            "RG-50.549" not in file):

            # convert file to docx, storing it in an untracked folder called temp
            file_docx = file + 'x'
            command = 'textutil -convert docx ' + file + ' -output ' + 'non_core_temp/'+ file_docx 
            call(command, shell=True)

            # append to the array
            core_doc_asset.append(file_docx)

    # go to created dir
    os.chdir('non_core_temp')

    # get the units for each file, store them and update tracker
    for file in core_doc_asset:
        # get text units for this entry
        units = getTextUnits(file)

        if units:
            # get RG number
            rg_number = file.split("_")[0]

            # find last occurrence of '.' and replace it with '*' 
            k = rg_number.rfind(".")
            mongo_rg = rg_number[:k] + "*" + rg_number[k+1:]

            # insert units on the output collection
            h.update_field(DB, OUTPUT, "shelfmark", mongo_rg, "structured_transcript", units)

            original_filename = file[:-1]
            # update status on the stracker
            h.update_field(DB, TRACKER, "microsoft_doc_file", original_filename, "status", "Processed")
            h.update_field(DB, TRACKER, "microsoft_doc_file", original_filename, "extraction_method", "transcribe_non_core_doc")
        else:
            print(file)
            missing_count += 1

    
    # success
    pprint.pprint("Non-core doc files were successfully processed, but there  " +  str(missing_count) + " missing")

if __name__ == "__main__":
    createStructuredTranscriptDocx()
    #TODO account for the fact that ongoing answers in different paragaphs can exist
   