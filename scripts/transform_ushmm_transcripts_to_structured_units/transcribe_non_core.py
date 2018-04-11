import sys, glob, os
helper_path = os.path.join("..", "..", "utils")
sys.path.insert(0, helper_path)
import helper_mongo as h

os.chdir("../../data/")
from docx import Document
from collections import defaultdict
from subprocess import call

import pprint
import constants
import re

TRACKER = constants.TRACKER_COLLECTION
OUTPUT = constants.OUTPUT_COLLECTION
DB = constants.DB

def safePrint(str_):
    return ''.join(i for i in str_ if ord(i) < 128)

def getMonologue(filename):
    doc = Document(filename)
    
    monologue = ""
    
    # ignore all the -TITLE
    # iterate over all paragraphs to get text units
    for para in doc.paragraphs:
        paragraph = para.text
        
        # ensure it is not an empty line
        if paragraph:
            
            # ensure it is not a header type
            if paragraph.count('-') < 2 and paragraph.count('.') < 2:
                monologue += paragraph + ' '
    
    return list({'unit': monologue})

def getUnstructuredUnits(filename):
    doc = Document(filename)
    units = list()
    isHeader = True
    for para in doc.paragraphs:
        paragraph = para.text
        
        # ensure it is not an empty line
        if paragraph:
            # ignore the initial header info
            if isHeader:
                if paragraph.split() > 5:
                    units.append(paragraph)
                    isHeader = False

            else:
                units.append({'unit': paragraph})

    return units   


def get405Units(filename):
    print(filename)
    doc = Document(filename)
    units = list()
    isHeader = True
    for para in doc.paragraphs:
        paragraph = para.text
        
        # ensure it is not an empty line
        if paragraph:
            # check if there is a question
            para_partition = paragraph.split(')', 1)
            
            
            # it is a Q & A
            if isHeader:
                # check if we found the start of the interview - first question
                if (len(para_partition) > 1 and 
                    para_partition[0][len(para_partition[0]) - 1] == '?'):
                    print(para_partition)
                    # ignore the initial ? for the question
                    units.append({'unit': para_partition[0][1:]})
                    units.append({'unit:': para_partition[1]})
                    isHeader = False
            else:
                if (len(para_partition) > 1 and 
                    para_partition[0][len(para_partition[0]) - 1] == '?'):
                    # ignore the initial ? for the question
                    units.append({'unit': para_partition[0][1:]})
                    units.append({'unit:': para_partition[1]})
                else:
                    units.append({'unit': paragraph})

    # in case it is a monologue
    #if not units:
        #for para in doc.paragraphs:
    if ".0031" in filename:
        pprint.pprint(units)
    return units

def getTextUnits(filename):
    doc = Document(filename)
    units = list()
    
    unit_tracker = defaultdict(int)
    
    non_units = ["name:", "date:", "thesis:", "currently:", "note", "comment", "grandparents:", "transcript:", "note:"]

    # iterate over all paragraphs to get text units
    for para in doc.paragraphs:
        paragraph = para.text
        
        # ensure it is not an empty line
        if paragraph:
            # get first word
            unit_type = paragraph.partition(' ')[0]
            # in case it is in the format of e.g 'George Salton:'
            
            # e.g AJ:, B.
            m = re.compile('[A-Z][A-Z]?[.|:|-]')
            type1 = m.match(unit_type)

            # e.g [WJ]
            n = re.compile('\[[A-Z][A-Z]\]')
            typer2= n.match(unit_type)

            # for interviews that have interviewee name at the beginning of unit
            #if unit_type.isupper() and unit_type.isalpha() and len(unit_type) > 1:

            # else parse them according to formatting guidelines
            if ("Question:" in unit_type or
                type1 or
                "Answer:" in unit_type or 
                typer2):

                safePrint(unit_type)
                units.append({'unit': paragraph})
                # update tracker
                unit_tracker[unit_type] += 1

            elif (unit_type.endswith(':') and
                    unit_type.lower() not in non_units and
                    unit_type[:-1].isalpha()):
                
                safePrint(unit_type)
                units.append({'unit': paragraph})
                # update tracker
                unit_tracker[unit_type] += 1
            
            # backup method,in case it is in the format of e.g 'George Salton:'
            elif len(paragraph.split()) > 3:
                backup_type = paragraph.split()[1]
                backup_two = paragraph.split()[2]

                safePrint(backup_type)

                if ((':' in backup_type and backup_type.lower() not in non_units) or
                    (':' in backup_two and backup_two.lower() not in non_units)): 
                    units.append({'unit': paragraph})
                    # update tracker
                    unit_tracker[unit_type] += 1
            # "Marion Muscati:" or "Vicente del Bosque"
            '''elif len(paragraph.split()) > 3:
                backup_type = paragraph.split()[2]
                third_backup = paragraph.split()[3]
                
                if filename == "RG-50.470.0012_trs_en.docx":
                    print third_backup
                if ((':' in backup_type and backup_type.lower() not in non_units) or 
                    (':' in third_backup and third_backup not in non_units)): 
                    units.append({'unit': paragraph})
                    # update tracker
                    unit_tracker[unit_type] += 1'''
 

    # check if backup method needed
    if len(unit_tracker) < 2:
        # the remaining 50.062 series is composed of monologues
        if "RG-50.062" in filename:
            units = getMonologue(filename)
            return units
        
        elif "RG-50.233" in filename:
            units = getUnstructuredUnits(filename)

        elif "RG-50.405" in filename:
            units = get405Units(filename)
        
        else:
            return []

    return units

def createStructuredTranscriptDocX():
    docx_assets = []
    count = 0
    for file in glob.glob("*.docx"):
         # RG numbers for the non-core asset
        if ("RG-50.030" not in file and
            "RG-50.106" not in file and
            "RG-50.549" not in file):
            docx_assets.append(file)

    # get the units for each file, store them and update tracker
    for file in docx_assets:
        # get text units for this entry
        units = getTextUnits(file)

        if units:
            # get RG number
            rg_number = file.split("_")[0]

            # find last occurrence of '.' and replace it with '*' 
            k = rg_number.rfind(".")
            mongo_rg = rg_number[:k] + "*" + rg_number[k+1:]

            # insert units on the output collection
            h.update_field(DB, OUTPUT, "shelfmark", mongo_rg, "structured_transcript", units)

            # update status on the stracker
            h.update_field(DB, TRACKER, "microsoft_doc_file", file, "status", "Processed")
        else:
            count += 1
            print(file)
    print count

    # success
    pprint.pprint("Core_docx_asset was successfully processed.")
    

if __name__ == "__main__":
    createStructuredTranscriptDocX()