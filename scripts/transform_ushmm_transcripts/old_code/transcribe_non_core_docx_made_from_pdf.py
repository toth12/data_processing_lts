import sys, glob, os
import helper_mongo as h
from docx import Document
from subprocess import call
from data_spec import create_dictionary_of_file_list
from get_text_units import getTextUnits

import pprint
import constants
import re
import pdb

TRACKER = constants.USHMM_TRACKER_COLLECTION
OUTPUT = constants.OUTPUT_COLLECTION_USHMM
DB = constants.DB
INPUT_FOLDER=constants.INPUT_FOLDER_USHMM_TRANSCRIPTS_PDF_TRANSFORMED_TO_DOCS
OUTPUT_FOLDER_USHMM_PROCESSING_LOGS=constants.OUTPUT_FOLDER_USHMM_PROCESSING_LOGS 



def getUnstructured50_583Units(filename):
    """
    Returns the unstructured units of the RG-50.583 series
    These interviews did not Question and Answer to distinguish units but uses the distinction of bold and italic
    """
    doc = Document(filename)
    units = list()

    
    

    # iterate over all paragraphs to get text units
    for para in doc.paragraphs:
        paragraph = para.text
        
        # ensure paragraph is not just empty line
        hasText = paragraph.lstrip()
        # ensure it is not an empty line
        if hasText:


            # marks the end of the interview
            if "shttp://collections.ushmm.org/search/catalog/irn43329" in paragraph.lower() in paragraph.lower() :
                break
            else:
                units.append({'unit':paragraph})
    return units

def getUnstructured_50_462_0005_Units(filename):
    """
    Returns the units for the RG-50.616 interviews, it does not try to find questions and answers;
    units are separated by empty lines in the original file
    
    """
    doc = Document(filename)
    units = list()
    
    
    for para in doc.paragraphs:
        paragraph = para.text
        if len(paragraph.strip())!=0:
            units.append({'unit': paragraph})
    
    return units


def getUnstructured_50_462_0032_Units(filename):
    """
    Returns the units for the RG-50.616 interviews, it does not try to find questions and answers;
    units are separated by empty lines in the original file
    
    """
    doc = Document(filename)
    units = list()
    
    
    for para in doc.paragraphs:
        paragraph = para.text
        if len(paragraph.strip())!=0:
            if ((paragraph[0:2]=='LC') or (paragraph[0:2]=='MS')):
                units.append({'unit': paragraph})
            elif len(units)>0:
                units[-1]['unit']=units[-1]['unit']+ ' '+paragraph
            else:
                units.append({'unit': paragraph})
    return units


def getTextUnits_old(filename):
    doc = Document(filename)
    units = list()
    # iterate over all paragraphs to get text units
    for para in doc.paragraphs:
        paragraph = para.text
        
        # ensure it is not an empty line
        if len(paragraph.strip())>0:
            # get first word
            unit_type = paragraph.partition(' ')[0]
            
            # e.g [d]
            m = re.compile('[A-Z][.|:]')
            type1 = m.match(unit_type)

            # e.g [WJ]
            n = re.compile('\[[A-Z][A-Z]\]')
            typer2= n.match(unit_type)

            # else parse them according to formatting guidelines
            if ("Question:" in unit_type or
                type1 or
                "Answer:" in unit_type or 
                typer2):

                units.append({'unit': paragraph})
            
            # backup for 2 interviews that don't match any of the patterns above
            elif (filename == "RG-50.030.0336_trs_en.docx" or 
                filename == "RG-50.030.0335_trs_en.docx"):

                if ("Interviewer:" in unit_type or 
                    "Theodore:" in unit_type or 
                    "MR." in unit_type):
                    
                    units.append({'unit': paragraph})
            else:
                if len(units)>0:
                    units[-1]['unit']= units[-1]['unit']+paragraph 
                else:   
                    units.append({'unit': paragraph})
    return units



def create_dictionary_of_file_list(filelist):
    result={}
    for file in filelist:
        original_filename = file.split('/')[-1]
        rg_number=original_filename.split("_")[0]

        # find last occurrence of '.' and replace it with '*' 
        k = rg_number.rfind(".")
        mongo_rg = rg_number[:k] + "*" + rg_number[k+1:]

        #check if already in the list

        if mongo_rg not in result.keys():
            result[mongo_rg]=[file]
        else:
            
            result[mongo_rg].append(file)

    
    return result


def createStructuredTranscriptDoc():
    """
    Processes the 509 doc files beloging to the core asset in data
    Core asset is identified by numbers RG-50.030, RG-50.106, RG-50.549
    """
    

    core_doc_asset = []
    missing_files=[]

    # get all the docx files that are part of the core asset
    for file in glob.glob(INPUT_FOLDER+"*.*"):
        # RG numbers for the core asset
        
        

        if not ("RG-50.030" in file or
            #this is questionable
            "RG-50.106" in file or
            "RG-50.549" in file):

            
            
            # append to the array
            core_doc_asset.append(file)

    
    core_doc_asset=create_dictionary_of_file_list(core_doc_asset)
    # get the units for each file, store them and update tracker
    not_processed=0
    processed_doc=0
    
    for mongo_rg in core_doc_asset:
        # get text units for this entry
        processed=[]
        result=[]
        
        for file in core_doc_asset[mongo_rg]:

            
            units = getTextUnits(file)
            
            if units:
                #replace white spaces
                for i,element in enumerate(units):
                    units[i]['unit']=' '.join(element['unit'].split())
                result.extend(units)
            
                processed.append(True)
            else:
                #check if processed
                processed.append(False)
        #set the method used to transform the transcript
        h.update_field(DB, TRACKER, "rg_number", mongo_rg, "method", "transcribe_non_core_docx_made_from_pdf")

        not_processed=not_processed+1
        if False in processed:
            h.update_field(DB, TRACKER, "rg_number", mongo_rg, "status", "Unprocessed")
            missing_files.append(' '.join(core_doc_asset[mongo_rg]))

            not_processed=not_processed+1
        else:
            # insert units on the output collection
            h.update_field(DB, OUTPUT, "shelfmark", 'USHMM '+mongo_rg, "structured_transcript", result)

                
            # update status on the stracker
                
            h.update_field(DB, TRACKER, "rg_number", mongo_rg, "status", "Processed")
            processed_doc=processed_doc+1

    
    #write the missing files to text file
    file = open(OUTPUT_FOLDER_USHMM_PROCESSING_LOGS+'transcribe_non_core_docx_made_from_pdf_failed.txt','w')
    file.write('\n'.join(missing_files))
    # success
    pprint.pprint("Core_doc_asset was successfully processed.")

if __name__ == "__main__":
    createStructuredTranscriptDoc()

   

    #getTextUnits()
    """
    result = h.query(DB, OUTPUT, { "structured_transcript": {'$exists': 'true'}}, {'id': 0})
    pprint.pprint(result)
    """
